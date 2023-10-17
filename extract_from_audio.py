import whisper
import os
import sys
import cv2
import numpy as np
import glob
from PIL import Image
from pydub import AudioSegment
from alive_progress import alive_bar
import docx 
from docx.shared import Pt
from docx.shared import Inches


def extract_text(file):   
    #convert video to audio 
    mp4_version = AudioSegment.from_file(file, "mp4")
    mp4_version.export(out_f = "Temp/audio.mp4", format = "wav")
    #os.system("ffmpeg -i "file" -map 0:a -acodec copy Temp/audio.mp4")
    model = whisper.load_model("base")
    result = model.transcribe("Temp/audio.mp4", verbose = False)
    return result["segments"]

def mse(img1, img2):
   h, w = img1.shape
   diff = cv2.subtract(img1, img2)
   err = np.sum(diff**2)
   mse = err/(float(h*w))
   return mse

def extract_images(file):
    #clean
    
    os.system('rm Temp/Data/*.jpg')

    cap=cv2.VideoCapture(file)
    fps = cap.get(cv2.CAP_PROP_FPS)
    totalNoFrames = cap.get(cv2.CAP_PROP_FRAME_COUNT)

    images=[]
    times=[]
    pwd=os.getcwd()
    # Read the video from specified path
    cam = cv2.VideoCapture(file)
    
    try:
        # creating a folder named data
        if not os.path.exists('Temp/data'):
            os.makedirs('Temp/data')
    
    # if not created then raise error
    except OSError:
        print ('Error: Creating directory of data')
    
    # frame
    #inti time
    currentframe = 0
    res=0
    with alive_bar(int(totalNoFrames),force_tty=True) as bar:
        while(True):
        
            # reading from frame
            ret,frame = cam.read()

            if ret:
                # if video is still left continue creating images
                time=float(currentframe/fps)
                #print(durationInSeconds-time)
                name = './Temp/data/frame' + str(currentframe) + '.jpg'
                path =pwd+'/Temp/data/frame' + str(currentframe) + '.jpg'    
                if(currentframe>0):
                    img1 = cv2.cvtColor(frame_old, cv2.COLOR_BGR2GRAY)
                    img2 = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                    res=mse(img1,img2)
                    if (res>=4.0):
                        #print ('Creating...' + name+" "+str(res))
                        # writing the extracted images
                        cv2.imwrite(name, frame)
                        times.append(time)
                        images.append(path)           

                else:
                    #print ('Creating...' + name+" "+str(res))
                    # writing the extracted images
                    cv2.imwrite(name, frame)
                    times.append(time)
                    images.append(path)

                currentframe += 1
                frame_old=frame
            else:
                break
            bar()
    
    # Release all space and windows once done
    cam.release()
    cv2.destroyAllWindows()
    return times,images


pathinput= "Videos/" 

for name in  os.listdir(pathinput):
    file="Videos/"+name
    print("Using Video : "+name)
    print("Extract_images \n")
    array=extract_images(file)
    print("Extract_text \n")
    data=extract_text(file)
    end_old=0.0
    temo="Output/"+name+".docx"
    doc = docx.Document()
    doc.add_heading(name, 0)

    print("Autput to Docx \n")

    with alive_bar(int(len(data)),force_tty=True) as bar:
        for i, seg in enumerate(data):
            doc.add_paragraph(str(i+1) +"- "+seg['text']) 
            count=0    
            for time in array[0]:
                if time >=end_old and time <= seg['end']:
                    doc.add_picture(array[1][count],width=Inches(6.9))
                count+=1
            end_old=seg['end']
            bar()

    doc.save(temo)
