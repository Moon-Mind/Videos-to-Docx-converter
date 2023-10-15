# Import docx NOT python-docx 
import pwd
import docx 
from docx.shared import Pt 

# Create an instance of a word document 
doc = docx.Document() 


# Adding paragraph with Increased font size 
doc.add_paragraph('GeeksforGeeks is a Computer Science portal for geeks.') 

doc.add_picture('/home/moonmind/Development/extract_data_from_video/Temp/data/frame0.jpg')


# Now save the document to a location 
doc.save('gfg.docx')
